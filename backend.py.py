# Import necessary libraries
from flask import Flask, request, jsonify, send_file
from transformers import pipeline
from docx import Document
import os

# Initialize Flask app
app = Flask(__name__)

# Load an NLP model for text classification
nlp_model = pipeline("zero-shot-classification", model="facebook/bart-large-mnli")

# In-memory storage for requirements (replace with a database in production)
requirements = {
    "functional": [],
    "non_functional": [],
}

# Function to categorize requirements
def categorize_requirement(requirement):
    categories = ["functional", "non-functional"]
    result = nlp_model(requirement, candidate_labels=categories)
    return result['labels'][0]  # Return the most likely categorypython srs_backend.py

# Function to generate the SRS document
def generate_srs(filename="srs_document.docx"):
    doc = Document()
    doc.add_heading("Software Requirements Specification (SRS)", level=1)

    # Add sections for functional and non-functional requirements
    doc.add_heading("Functional Requirements", level=2)
    for req in requirements["functional"]:
        doc.add_paragraph(f"- {req}", style="List Bullet")

    doc.add_heading("Non-Functional Requirements", level=2)
    for req in requirements["non_functional"]:
        doc.add_paragraph(f"- {req}", style="List Bullet")

    # Save the document
    doc.save(filename)
    return filename

# API to add a requirement
@app.route("/add-requirement", methods=["POST"])
def add_requirement():
    data = request.json
    requirement = data.get("requirement")

    if not requirement:
        return jsonify({"error": "Requirement is required"}), 400

    # Categorize the requirement using AI
    category = categorize_requirement(requirement)

    # Add the requirement to the appropriate list
    if category == "functional":
        requirements["functional"].append(requirement)
    else:
        requirements["non_functional"].append(requirement)

    return jsonify({"message": "Requirement added successfully", "category": category})

# API to generate and download the SRS document
@app.route("/generate-srs", methods=["GET"])
def download_srs():
    try:
        # Generate the SRS document
        filename = generate_srs()

        # Send the file as a response
        return send_file(filename, as_attachment=True)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# API to list all requirements
@app.route("/list-requirements", methods=["GET"])
def list_requirements():
    return jsonify(requirements)

# API to clear all requirements
@app.route("/clear-requirements", methods=["DELETE"])
def clear_requirements():
    requirements["functional"].clear()
    requirements["non_functional"].clear()
    return jsonify({"message": "All requirements cleared"})

# Run the Flask app
if __name__ == "__main__":
    app.run(debug=True)